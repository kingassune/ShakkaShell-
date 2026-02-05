"""Report generator for multiple output formats.

Generates penetration testing reports in Markdown, HTML, DOCX, and PDF formats.
"""

import json
import logging
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Optional, List, Dict, Any

from shakka.reports.models import (
    Report,
    Finding,
    Evidence,
    Severity,
    ReportMetadata,
    CVSSScore,
    EvidenceType,
)
from shakka.reports.templates import (
    ReportTemplate,
    TemplateRegistry,
    TemplateRenderer,
    DEFAULT_MARKDOWN_TEMPLATE,
    DEFAULT_HTML_TEMPLATE,
)


logger = logging.getLogger(__name__)


class OutputFormat(Enum):
    """Supported output formats."""
    
    MARKDOWN = "markdown"
    HTML = "html"
    DOCX = "docx"
    PDF = "pdf"
    JSON = "json"
    
    @property
    def extension(self) -> str:
        """Get file extension for format."""
        extensions = {
            "markdown": ".md",
            "html": ".html",
            "docx": ".docx",
            "pdf": ".pdf",
            "json": ".json",
        }
        return extensions[self.value]


@dataclass
class GeneratorConfig:
    """Configuration for report generator."""
    
    default_format: OutputFormat = OutputFormat.MARKDOWN
    default_template: str = "default"
    include_evidence: bool = True
    include_executive_summary: bool = True
    auto_generate_summary: bool = True
    output_directory: str = "./reports"
    company_logo: Optional[str] = None
    custom_css: Optional[str] = None
    page_size: str = "A4"  # For PDF
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "default_format": self.default_format.value,
            "default_template": self.default_template,
            "include_evidence": self.include_evidence,
            "include_executive_summary": self.include_executive_summary,
            "auto_generate_summary": self.auto_generate_summary,
            "output_directory": self.output_directory,
            "company_logo": self.company_logo,
            "custom_css": self.custom_css,
            "page_size": self.page_size,
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "GeneratorConfig":
        """Create from dictionary."""
        return cls(
            default_format=OutputFormat(data.get("default_format", "markdown")),
            default_template=data.get("default_template", "default"),
            include_evidence=data.get("include_evidence", True),
            include_executive_summary=data.get("include_executive_summary", True),
            auto_generate_summary=data.get("auto_generate_summary", True),
            output_directory=data.get("output_directory", "./reports"),
            company_logo=data.get("company_logo"),
            custom_css=data.get("custom_css"),
            page_size=data.get("page_size", "A4"),
        )


class ReportGenerator:
    """Generates penetration testing reports in multiple formats."""
    
    def __init__(
        self,
        config: Optional[GeneratorConfig] = None,
        template_registry: Optional[TemplateRegistry] = None,
    ):
        """Initialize the report generator.
        
        Args:
            config: Generator configuration
            template_registry: Custom template registry
        """
        self.config = config or GeneratorConfig()
        self.templates = template_registry or TemplateRegistry()
        self._generated_reports: List[str] = []
    
    @property
    def generated_reports(self) -> List[str]:
        """Get list of generated report paths."""
        return self._generated_reports.copy()
    
    def generate(
        self,
        report: Report,
        output_format: Optional[OutputFormat] = None,
        output_path: Optional[str] = None,
        template_name: Optional[str] = None,
    ) -> str:
        """Generate a report in the specified format.
        
        Args:
            report: Report to generate
            output_format: Output format (default from config)
            output_path: Output file path (auto-generated if None)
            template_name: Template to use (default from config)
            
        Returns:
            Path to generated report
        """
        fmt = output_format or self.config.default_format
        template = self._get_template(template_name, fmt)
        
        # Auto-generate executive summary if enabled
        if self.config.auto_generate_summary and not report.metadata.executive_summary:
            report.metadata.executive_summary = self._generate_executive_summary(report)
        
        # Generate content based on format
        if fmt == OutputFormat.MARKDOWN:
            content = self._generate_markdown(report, template)
        elif fmt == OutputFormat.HTML:
            content = self._generate_html(report, template)
        elif fmt == OutputFormat.JSON:
            content = self._generate_json(report)
        elif fmt == OutputFormat.DOCX:
            content = self._generate_docx(report, template)
        elif fmt == OutputFormat.PDF:
            content = self._generate_pdf(report, template)
        else:
            raise ValueError(f"Unsupported format: {fmt}")
        
        # Determine output path
        if not output_path:
            output_path = self._generate_output_path(report, fmt)
        
        # Write to file
        self._write_output(output_path, content, fmt)
        self._generated_reports.append(output_path)
        
        logger.info(f"Generated report: {output_path}")
        return output_path
    
    def generate_all_formats(
        self,
        report: Report,
        output_directory: Optional[str] = None,
    ) -> Dict[OutputFormat, str]:
        """Generate report in all supported formats.
        
        Args:
            report: Report to generate
            output_directory: Output directory
            
        Returns:
            Dictionary mapping format to output path
        """
        out_dir = output_directory or self.config.output_directory
        paths = {}
        
        for fmt in [OutputFormat.MARKDOWN, OutputFormat.HTML, OutputFormat.JSON]:
            base_name = self._generate_base_name(report)
            output_path = str(Path(out_dir) / f"{base_name}{fmt.extension}")
            
            try:
                paths[fmt] = self.generate(report, fmt, output_path)
            except Exception as e:
                logger.error(f"Failed to generate {fmt.value}: {e}")
        
        return paths
    
    def preview(
        self,
        report: Report,
        output_format: OutputFormat = OutputFormat.MARKDOWN,
        template_name: Optional[str] = None,
    ) -> str:
        """Preview report without saving to file.
        
        Args:
            report: Report to preview
            output_format: Output format
            template_name: Template to use
            
        Returns:
            Generated content string
        """
        template = self._get_template(template_name, output_format)
        
        if output_format == OutputFormat.MARKDOWN:
            return self._generate_markdown(report, template)
        elif output_format == OutputFormat.HTML:
            return self._generate_html(report, template)
        elif output_format == OutputFormat.JSON:
            return self._generate_json(report)
        else:
            return self._generate_markdown(report, template)
    
    # =========================================================================
    # Format-specific generators
    # =========================================================================
    
    def _generate_markdown(self, report: Report, template: ReportTemplate) -> str:
        """Generate Markdown report."""
        renderer = TemplateRenderer(template)
        return renderer.render(report)
    
    def _generate_html(self, report: Report, template: ReportTemplate) -> str:
        """Generate HTML report."""
        # Use HTML template if available, otherwise convert from markdown
        if template.format_type == "html":
            renderer = TemplateRenderer(template)
            return renderer.render(report)
        
        # Convert markdown to HTML
        md_content = self._generate_markdown(report, template)
        return self._markdown_to_html(md_content, report)
    
    def _generate_json(self, report: Report) -> str:
        """Generate JSON report."""
        return json.dumps(report.to_dict(), indent=2, default=str)
    
    def _generate_docx(self, report: Report, template: ReportTemplate) -> bytes:
        """Generate DOCX report.
        
        Note: Requires python-docx package for full functionality.
        Returns a simple text representation if package not available.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(report.metadata.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Document info
            doc.add_paragraph(f"Classification: {report.metadata.classification}")
            doc.add_paragraph(f"Version: {report.metadata.version}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
            
            # Info table
            doc.add_heading("Document Information", level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = "Table Grid"
            
            info = [
                ("Client", report.metadata.client),
                ("Assessor", report.metadata.assessor),
                ("Assessment Type", report.metadata.assessment_type),
                ("Period", f"{report.metadata.start_date} - {report.metadata.end_date}"),
            ]
            for i, (key, value) in enumerate(info):
                table.rows[i].cells[0].text = key
                table.rows[i].cells[1].text = str(value) if value else "N/A"
            
            # Executive Summary
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(report.metadata.executive_summary or "No summary provided.")
            
            # Risk Overview
            doc.add_heading("Risk Overview", level=2)
            risk_table = doc.add_table(rows=6, cols=2)
            risk_table.style = "Table Grid"
            
            risk_data = [
                ("Critical", str(report.critical_count)),
                ("High", str(report.high_count)),
                ("Medium", str(report.medium_count)),
                ("Low", str(report.low_count)),
                ("Informational", str(report.info_count)),
                ("Total", str(report.total_findings)),
            ]
            for i, (sev, count) in enumerate(risk_data):
                risk_table.rows[i].cells[0].text = sev
                risk_table.rows[i].cells[1].text = count
            
            # Findings
            doc.add_heading("Findings", level=1)
            for i, finding in enumerate(report.sorted_findings, 1):
                doc.add_heading(f"{i}. {finding.title}", level=2)
                doc.add_paragraph(f"Severity: {finding.severity.value.title()}")
                doc.add_paragraph(f"CVSS Score: {finding.cvss_score:.1f}")
                doc.add_paragraph(f"Affected Asset: {finding.affected_asset or 'N/A'}")
                
                doc.add_heading("Description", level=3)
                doc.add_paragraph(finding.description)
                
                if finding.remediation:
                    doc.add_heading("Remediation", level=3)
                    doc.add_paragraph(finding.remediation)
            
            # Save to bytes
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
            
        except ImportError:
            logger.warning("python-docx not installed, returning markdown")
            return self._generate_markdown(report, template).encode()
    
    def _generate_pdf(self, report: Report, template: ReportTemplate) -> bytes:
        """Generate PDF report.
        
        Note: Requires weasyprint or similar package for full functionality.
        Returns HTML content if package not available.
        """
        try:
            from weasyprint import HTML
            
            html_content = self._generate_html(report, template)
            pdf_bytes = HTML(string=html_content).write_pdf()
            return pdf_bytes
            
        except ImportError:
            logger.warning("weasyprint not installed, returning HTML")
            return self._generate_html(report, template).encode()
    
    # =========================================================================
    # Helper methods
    # =========================================================================
    
    def _get_template(
        self, template_name: Optional[str], fmt: OutputFormat
    ) -> ReportTemplate:
        """Get appropriate template for format."""
        name = template_name or self.config.default_template
        
        # Try to get named template
        template = self.templates.get(name)
        
        # If format is HTML, prefer HTML template
        if fmt == OutputFormat.HTML and not template:
            template = self.templates.get("html")
        
        # Fallback to default
        if not template:
            template = self.templates.get("default") or ReportTemplate(
                name="fallback",
                report_template=DEFAULT_MARKDOWN_TEMPLATE,
            )
        
        return template
    
    def _generate_output_path(self, report: Report, fmt: OutputFormat) -> str:
        """Generate output file path."""
        base_name = self._generate_base_name(report)
        output_dir = Path(self.config.output_directory)
        output_dir.mkdir(parents=True, exist_ok=True)
        return str(output_dir / f"{base_name}{fmt.extension}")
    
    def _generate_base_name(self, report: Report) -> str:
        """Generate base filename for report."""
        date_str = datetime.now().strftime("%Y%m%d")
        client = report.metadata.client.replace(" ", "_")[:20] if report.metadata.client else "report"
        return f"{client}_{date_str}_{report.id[:8]}"
    
    def _write_output(self, path: str, content: Any, fmt: OutputFormat):
        """Write content to output file."""
        mode = "wb" if fmt in [OutputFormat.DOCX, OutputFormat.PDF] else "w"
        
        # Ensure directory exists
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        
        with open(path, mode) as f:
            if isinstance(content, bytes):
                f.write(content)
            else:
                f.write(content)
    
    def _generate_executive_summary(self, report: Report) -> str:
        """Auto-generate executive summary from findings."""
        if not report.findings:
            return "No security vulnerabilities were identified during this assessment."
        
        # Build summary
        parts = []
        
        # Overview
        parts.append(
            f"During this assessment, a total of {report.total_findings} security "
            f"findings were identified across the in-scope assets."
        )
        
        # Critical/High issues
        critical_high = report.critical_count + report.high_count
        if critical_high > 0:
            parts.append(
                f"Of these, {critical_high} are considered critical or high severity "
                "and require immediate attention."
            )
        
        # Risk level
        if report.risk_score >= 70:
            parts.append(
                "The overall security posture is considered HIGH RISK and "
                "significant remediation efforts are recommended before production deployment."
            )
        elif report.risk_score >= 40:
            parts.append(
                "The overall security posture is considered MODERATE RISK. "
                "Remediation of identified issues is recommended in a timely manner."
            )
        else:
            parts.append(
                "The overall security posture is considered LOW RISK, though "
                "addressing identified issues would further strengthen the security posture."
            )
        
        # Top findings
        if report.critical_count > 0:
            critical = report.get_findings_by_severity(Severity.CRITICAL)
            parts.append(
                f"The most critical finding is '{critical[0].title}' which affects "
                f"{critical[0].affected_asset or 'the target system'}."
            )
        
        return " ".join(parts)
    
    def _markdown_to_html(self, markdown: str, report: Report) -> str:
        """Convert markdown to HTML with styling."""
        try:
            import markdown as md
            html_body = md.markdown(markdown, extensions=["tables", "fenced_code"])
        except ImportError:
            # Basic conversion without markdown package
            html_body = f"<pre>{markdown}</pre>"
        
        # Wrap in HTML template
        return f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{report.metadata.title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; }}
        h1, h2, h3 {{ color: #2c3e50; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background: #f5f5f5; }}
        pre {{ background: #f5f5f5; padding: 15px; overflow-x: auto; }}
        code {{ background: #f5f5f5; padding: 2px 5px; }}
        {self.config.custom_css or ''}
    </style>
</head>
<body>
{html_body}
</body>
</html>'''


# =============================================================================
# Convenience functions
# =============================================================================

def create_report(
    title: str = "Penetration Testing Report",
    client: str = "",
    assessor: str = "",
) -> Report:
    """Create a new report with basic metadata.
    
    Args:
        title: Report title
        client: Client name
        assessor: Assessor name
        
    Returns:
        New Report instance
    """
    metadata = ReportMetadata(
        title=title,
        client=client,
        assessor=assessor,
        start_date=datetime.now(),
    )
    return Report(metadata=metadata)


def create_finding(
    title: str,
    description: str,
    severity: Severity,
    cvss_score: Optional[float] = None,
    affected_asset: str = "",
    remediation: str = "",
) -> Finding:
    """Create a new finding.
    
    Args:
        title: Finding title
        description: Full description
        severity: Severity level
        cvss_score: CVSS score (optional)
        affected_asset: Affected system/asset
        remediation: Remediation guidance
        
    Returns:
        New Finding instance
    """
    cvss = CVSSScore(score=cvss_score) if cvss_score else None
    
    return Finding(
        title=title,
        description=description,
        severity=severity,
        cvss=cvss,
        affected_asset=affected_asset,
        remediation=remediation,
    )


def quick_generate(
    report: Report,
    output_format: str = "markdown",
    output_path: Optional[str] = None,
) -> str:
    """Quick report generation with sensible defaults.
    
    Args:
        report: Report to generate
        output_format: Format string (markdown, html, json, docx, pdf)
        output_path: Output path (optional)
        
    Returns:
        Path to generated report
    """
    fmt = OutputFormat(output_format)
    generator = ReportGenerator()
    return generator.generate(report, fmt, output_path)
